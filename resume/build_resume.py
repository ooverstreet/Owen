#!/usr/bin/env python3
"""Build Owen's resume as a Google Docs-style .docx and Times-style .pdf."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Twips
from reportlab.lib.colors import HexColor, black
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUT = Path(__file__).resolve().parent
BLUE = RGBColor(0x11, 0x55, 0xCC)
RL_BLUE = HexColor("#1155CC")
LINK = "https://www.linkedin.com/in/owen-overstreet-2119776"
EMAIL = "OwenOverstreet@ymail.com"
FONT_DIR = Path("/usr/share/fonts/truetype/liberation")

JOBS = [
    {
        "title": "Technical Analyst/Senior Business Analyst",
        "dates": "Feb 2022 - Present",
        "company": "ICE Mortgage Technology (Fka: Black Knight Inc.)",
        "bullets": [
            "Partner with clients, BA's and Developers to gather and document business and data requirements; translate needs into configuration, coding, and data changes and stay with each request through implementation.",
            "Write T-SQL to research production issues, validate data across environments, produce Ad hoc reports, and create scripts plus Change Orders / work orders for DBA execution.",
            "Perform QA testing with the QA team and client before Prod/UAT sign off; exercise the change, confirm data and configs, and stay with the request through release.",
            "Document the work end to end — Visio Flow diagrams, Technical Documents, SOW's, change/work orders, and User Stories — so business rules, data, and code are clear for sign off.",
            "Research Prod issues/incident tickets (via Service Now) with SQL, Visual Studio, and the Web App; identify solutions and advise Client Support of findings.",
            "Vibe code with Cursor and AI agents (plus VS Code) as a daily practice to write SQL, scripts, QA checks, and documentation faster while keeping production standards.",
        ],
    },
    {
        "title": "IT Business Systems Analyst III",
        "dates": "June 2012 – 2022",
        "company": "Black Knight Financial Services (Fka: LPSVCS/Lender Processing Services), Jacksonville, FL",
        "bullets": [
            "Dedicated support for Wells Fargo on LoanSphere BK/FC/FCR/IM and mappings for MSP (Mortgage Servicing Platform).",
            "Gathered client requirements and implemented configuration and data-mapping changes across six environments (UAT, QA1, QA2, DEV, PA, and PROD); verified data landed in the correct environment before release.",
            "Wrote T-SQL (Insert/Update/Delete) Mapping Scripts, Attorney Updates, and Dynamic Zip codes; submitted tickets for DBA's to Execute and tracked through all Environments.",
            "Wrote AdHoc's for Client/Team via SQL; queried different environments to validate Configurations and find data anomalies, including XML errors.",
            "Created and sent SOW's to Sr. Mgmt. for sign off before implementing to Prod; attached RFW's, SOW's, and Config Sheets to TFS / Service Desk tickets.",
            "Updated SharePoint with Current Configurations pulled from Production for Client and Teams.",
            "Ran QA / UAT testing with the Client (load loans for testing, config checks) before production sign off; mentored Junior associates and provided backup for Manager when out.",
        ],
    },
    {
        "title": "Product Support / Data Analyst",
        "dates": "February 2011 – June 2012",
        "company": "DMEautomotive, Jacksonville, FL",
        "bullets": [
            "Support for The Goodyear Tire and Rubber Co. to facilitate the rollout of their Point of Sale Incentive program across U.S and Canada; worked with POS vendors and Goodyear to get customer data into the application.",
            "Provide Ad-Hoc reports for Goodyear and their customers using T-SQL; load data into SQL Server and run stored procedures to add or remove records.",
            "Check data load jobs daily and SSIS tasks; trouble shoot failures with T-SQL and fix errors so loads complete as quick as possible.",
        ],
    },
    {
        "title": "Lead Support Developer / Analytical Development for LoanSphere",
        "dates": "March 2008 – February 2011",
        "company": "Lender Processing Services Default Solutions / LPSDS, Jacksonville, FL",
        "bullets": [
            "Use T-SQL daily to write scripts, return data for reports, and make changes to the Desktop App (Foreclosures, Bankruptcies) which supports thousands of end-users. Follow the process until completion.",
            "Create new Processes, Holds, Issues, Events, Data Forms and Close reasons for LoanSphere; document the change and contact clients to ensure the proper data has been given and business needs are met.",
            "QA-tested issues reported by clients to identify if the issue is a software defect or if the user needs training; provide Ad-Hoc reports via Excel and work tickets until complete.",
        ],
    },
    {
        "title": "IT / Desktop Support / Jr. Network Associate",
        "dates": "September 2006 – March 2008",
        "company": "Lender Processing Services / Fidelity Process Management, Jacksonville, FL",
        "bullets": [
            "Support 300+ users and Fidelity's Process Management application (attorney foreclosure/bankruptcy loans) with over 5,000 users; identify if a reported issue is a software defect or if the user needs training.",
            "Provide detailed specifications when reporting a software defect by entering a development ticket; document the defect and the specification needed for the enhancement.",
        ],
    },
]

SKILLS = [
    (
        "Data & SQL:",
        "SSMS (T-Sql); ad-hoc reporting; stored procedures; SSIS data-load jobs; data mapping; data validation; XML; Microsoft Excel, Word, Outlook, PowerPoint.",
    ),
    (
        "QA testing:",
        "UAT and lower-environment testing; Prod/UAT sign off; defect vs. training triage; test-loan loads; data-quality checks across DEV / QA / UAT / PROD.",
    ),
    (
        "Documenting:",
        "Requirements gathering; Visio Flow diagrams; Technical Documents; SOW's / RFW's / work orders; Config Sheets; SharePoint; User Stories.",
    ),
    (
        "Vibe coding / AI:",
        "Cursor and AI agents for SQL, scripts, QA checks, and documentation; Visual Studio 2022; VS Code; Python (data extraction).",
    ),
    (
        "Platforms:",
        "Loan Sphere / MSP / Invoice Management (ICE Mortgage FKA Black Knight); Process Transfer Tool / Configuration Loading Tools / Loan Drop Tools / Mortgage Web Software; Service Now, Azure, TFS.",
    ),
]

SUMMARY = (
    "Mortgage-technology analyst with 15+ years using T-SQL to query, validate, map, and report on loan and servicing data. "
    "Partner with business stakeholders, product, and development to gather requirements, document the work, run QA / UAT testing, "
    "and deliver data and configuration changes through production. Active every day in QA testing, technical documentation, "
    "and vibe coding with Cursor and AI agents to write SQL, scripts, tests, and docs faster without dropping production standards."
)


def _set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _tight_paragraph(p, after=0, before=0, line=240):
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = Twips(line)


def build_docx(path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight_paragraph(p, after=0)
    r = p.add_run("OWEN OVERSTREET")
    _set_run_font(r, size=18, bold=True)

    for text, size, italic, bold in [
        ("Data Analyst  ·  Technical Analyst  ·  Business Analyst", 12, True, False),
        ("QA testing  ·  Technical documentation  ·  Vibe coding (Cursor / AI agents)", 11, False, True),
        ("Yulee, FL", 11, False, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tight_paragraph(p, after=0)
        r = p.add_run(text)
        _set_run_font(r, size=size, italic=italic, bold=bold)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight_paragraph(p, after=0)
    r = p.add_run("(P) 904-651-4256   (E) ")
    _set_run_font(r, size=11)
    r = p.add_run(EMAIL)
    _set_run_font(r, size=11, color=BLUE)
    r.font.underline = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight_paragraph(p, after=6)
    r = p.add_run("www.linkedin.com/in/owen-overstreet-2119776")
    _set_run_font(r, size=11, color=BLUE)
    r.font.underline = True

    def heading(text):
        p = doc.add_paragraph()
        _tight_paragraph(p, before=8, after=4)
        r = p.add_run(text)
        _set_run_font(r, size=12, bold=True)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    heading("PROFESSIONAL SUMMARY")
    p = doc.add_paragraph()
    _tight_paragraph(p, after=2)
    r = p.add_run(SUMMARY)
    _set_run_font(r, size=11)

    heading("TECHNICAL KNOWLEDGE")
    for label, body in SKILLS:
        p = doc.add_paragraph(style="List Bullet")
        _tight_paragraph(p, after=1)
        r = p.add_run(label + " ")
        _set_run_font(r, size=11, bold=True)
        r = p.add_run(body)
        _set_run_font(r, size=11)

    heading("PROFESSIONAL EXPERIENCE")
    for job in JOBS:
        p = doc.add_paragraph()
        _tight_paragraph(p, before=6, after=0)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run(job["title"])
        _set_run_font(r, size=11, bold=True)
        r = p.add_run("\t" + job["dates"])
        _set_run_font(r, size=11, bold=True)

        p = doc.add_paragraph()
        _tight_paragraph(p, after=1)
        r = p.add_run(job["company"])
        _set_run_font(r, size=11, italic=True)

        for b in job["bullets"]:
            p = doc.add_paragraph(style="List Bullet")
            _tight_paragraph(p, after=1)
            r = p.add_run(b)
            _set_run_font(r, size=11)

    heading("EDUCATION AND CERTIFICATIONS")
    edu = [
        "Database course — Florida State College at Jacksonville (formerly FCCJ South).",
        "Python for data extraction (in progress). Cursor and AI-agent development (vibe coding).",
        "MCP (Microsoft Certified Professional) — May 2000.  A+ — February 2006.  N+ — April 17, 2006.",
        "Network + Security Course, Jacksonville, FL.",
    ]
    for line in edu:
        p = doc.add_paragraph(style="List Bullet")
        _tight_paragraph(p, after=1)
        r = p.add_run(line)
        _set_run_font(r, size=11)

    doc.save(path)


def build_pdf(path: Path) -> None:
    pdfmetrics.registerFont(TTFont("TNR", str(FONT_DIR / "LiberationSerif-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("TNR-B", str(FONT_DIR / "LiberationSerif-Bold.ttf")))
    pdfmetrics.registerFont(TTFont("TNR-I", str(FONT_DIR / "LiberationSerif-Italic.ttf")))
    pdfmetrics.registerFont(TTFont("TNR-BI", str(FONT_DIR / "LiberationSerif-BoldItalic.ttf")))
    from reportlab.pdfbase.pdfmetrics import registerFontFamily

    registerFontFamily("TNR", normal="TNR", bold="TNR-B", italic="TNR-I", boldItalic="TNR-BI")

    styles = getSampleStyleSheet()
    center = ParagraphStyle(
        "center", parent=styles["Normal"], fontName="TNR", fontSize=11, leading=13, alignment=TA_CENTER, spaceAfter=0, spaceBefore=0
    )
    name = ParagraphStyle("name", parent=center, fontName="TNR-B", fontSize=18, leading=20, spaceAfter=2)
    roles = ParagraphStyle("roles", parent=center, fontName="TNR-I", fontSize=12, leading=14)
    skills_line = ParagraphStyle("skills_line", parent=center, fontName="TNR-B", fontSize=11, leading=13, spaceAfter=2)
    contact = ParagraphStyle("contact", parent=center, fontSize=11, leading=13)
    section = ParagraphStyle(
        "section",
        parent=styles["Normal"],
        fontName="TNR-B",
        fontSize=12,
        leading=14,
        spaceBefore=10,
        spaceAfter=4,
        borderPadding=0,
        textColor=black,
    )
    body = ParagraphStyle("body", parent=styles["Normal"], fontName="TNR", fontSize=11, leading=13, alignment=TA_LEFT, spaceAfter=2)
    bullet = ParagraphStyle("bullet", parent=body, leftIndent=14, bulletIndent=0, spaceAfter=1.5)
    job_title = ParagraphStyle("job_title", parent=body, fontName="TNR-B", spaceBefore=6, spaceAfter=0)
    company = ParagraphStyle("company", parent=body, fontName="TNR-I", spaceAfter=2)

    def section_table(title: str):
        t = Table([[Paragraph(title, section)]], colWidths=[7.0 * inch])
        t.setStyle(
            TableStyle(
                [
                    ("LINEBELOW", (0, 0), (-1, -1), 0.6, black),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]
            )
        )
        return t

    story = [
        Paragraph("OWEN OVERSTREET", name),
        Paragraph("Data Analyst  ·  Technical Analyst  ·  Business Analyst", roles),
        Paragraph("QA testing  ·  Technical documentation  ·  Vibe coding (Cursor / AI agents)", skills_line),
        Paragraph("Yulee, FL", contact),
        Paragraph(f'(P) 904-651-4256 &nbsp;&nbsp; (E) <font color="#1155CC"><u>{EMAIL}</u></font>', contact),
        Paragraph(f'<link href="{LINK}"><font color="#1155CC"><u>www.linkedin.com/in/owen-overstreet-2119776</u></font></link>', contact),
        Spacer(1, 6),
        section_table("PROFESSIONAL SUMMARY"),
        Paragraph(SUMMARY, body),
        section_table("TECHNICAL KNOWLEDGE"),
    ]
    for label, text in SKILLS:
        story.append(Paragraph(f"•  <b>{label}</b> {text}", bullet))

    story.append(section_table("PROFESSIONAL EXPERIENCE"))
    for job in JOBS:
        row = Table(
            [[
                Paragraph(job["title"], job_title),
                Paragraph(job["dates"], ParagraphStyle("dates", parent=job_title, alignment=TA_RIGHT)),
            ]],
            colWidths=[5.15 * inch, 1.85 * inch],
        )
        row.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]
            )
        )
        block = [row, Paragraph(job["company"], company)]
        for b in job["bullets"]:
            block.append(Paragraph(f"•  {b}", bullet))
        story.append(KeepTogether(block))

    story.append(section_table("EDUCATION AND CERTIFICATIONS"))
    for line in [
        "Database course — Florida State College at Jacksonville (formerly FCCJ South).",
        "Python for data extraction (in progress). Cursor and AI-agent development (vibe coding).",
        "MCP (Microsoft Certified Professional) — May 2000.  A+ — February 2006.  N+ — April 17, 2006.",
        "Network + Security Course, Jacksonville, FL.",
    ]:
        story.append(Paragraph(f"•  {line}", bullet))

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.5 * inch,
        title="OwenResume.A",
        author="Owen Overstreet",
    )
    doc.build(story)


if __name__ == "__main__":
    build_docx(OUT / "OwenResume.docx")
    build_pdf(OUT / "OwenResume.pdf")
    build_pdf(OUT / "Owen-Overstreet-Resume.pdf")
    print("wrote", OUT / "OwenResume.docx")
    print("wrote", OUT / "OwenResume.pdf")
    print("wrote", OUT / "Owen-Overstreet-Resume.pdf")
